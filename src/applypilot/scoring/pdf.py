"""Text-to-PDF/DOCX conversion for tailored resumes and cover letters.

Parses the structured text resume format, renders via an HTML/CSS template,
and exports to PDF using headless Chromium via Playwright, or to DOCX using
python-docx.

Supported formats: "pdf" (default), "docx".
"""

import logging
import re
from pathlib import Path
from typing import Any

from applypilot.config import TAILORED_DIR

# Valid document output formats
VALID_DOC_FORMATS = ("pdf", "docx")

log = logging.getLogger(__name__)


# ── URL / email auto-linking ─────────────────────────────────────────────

# Match http(s) URLs, bare domains (foo.com/path), and email addresses.
# Order matters: scheme'd URLs first, then emails, then bare domains.
_URL_RE = re.compile(
    r"""
    (
        # explicit scheme
        https?://[^\s<>()|]+
        |
        # email
        [A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}
        |
        # bare domain — at least one dot, common TLDs, optional path
        \b
        (?:[a-z0-9-]+\.)+(?:com|org|io|co|dev|app|ai|care|net|us|edu)
        (?:/[^\s<>()|]*)?
        \b
    )
    """,
    re.VERBOSE | re.IGNORECASE,
)


def _normalize_link(token: str) -> tuple[str, str]:
    """Return (display_text, href) for an auto-detected token.

    Emails get a mailto: prefix. Bare domains get an https:// prefix.
    The display text always preserves what the user wrote.
    """
    if "@" in token and "://" not in token:
        return token, f"mailto:{token}"
    if "://" not in token:
        return token, f"https://{token}"
    return token, token


def _split_text_with_links(text: str) -> list[tuple[str, str | None]]:
    """Split a string into (text_segment, href-or-None) tuples for rendering.

    Plain text yields (segment, None). Detected URLs/emails yield
    (display_text, href).
    """
    parts: list[tuple[str, str | None]] = []
    pos = 0
    for m in _URL_RE.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], None))
        token = m.group(0)
        # Trim trailing punctuation that the regex may have grabbed.
        trailing = ""
        while token and token[-1] in ".,;:!?)":
            trailing = token[-1] + trailing
            token = token[:-1]
        if token:
            display, href = _normalize_link(token)
            parts.append((display, href))
        if trailing:
            parts.append((trailing, None))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], None))
    return parts


# ── Resume Parser ────────────────────────────────────────────────────────

def parse_resume(text: str) -> dict:
    """Parse a structured text resume into sections.

    Supports two header layouts:

    1. **Pipeline-tailored** (compact): name → title → location → contact,
       all on consecutive non-blank lines, then a blank, then SUMMARY.

    2. **Master / ATS-friendly** (separated): a contact block (name + 1–2
       contact lines) → blank line → headline (job title / objective)
       → blank → SUMMARY.

    The presence of a blank line *inside* the pre-SUMMARY region is what
    separates the two — a blank break after the contact block means the
    headline is whatever comes after the break.

    Args:
        text: Full resume text.

    Returns:
        ``{"name": str, "title": str, "location": str, "contact": str,
           "sections": dict}``
    """
    lines = [line.rstrip() for line in text.strip().split("\n")]

    contact_block: list[str] = []
    headline_block: list[str] = []
    saw_break = False
    body_start = 0
    for i, line in enumerate(lines):
        if line.strip().upper() == "SUMMARY":
            body_start = i
            break
        if not line.strip():
            if contact_block:
                saw_break = True
            continue
        if saw_break:
            headline_block.append(line.strip())
        else:
            contact_block.append(line.strip())

    name = contact_block[0] if contact_block else ""

    # Title: the headline block when present (master format), otherwise
    # the second line of the contact block (compact format).
    if headline_block:
        title = " ".join(headline_block)
    elif len(contact_block) > 1:
        title = contact_block[1]
    else:
        title = ""

    # Location + contact extraction.
    location = ""
    contact = ""
    if headline_block:
        # Master format: contact_block has [name, "City, ST | phone | email", "LinkedIn: ... | GitHub: ..."]
        rest = contact_block[1:]
        if rest:
            first_rest = rest[0]
            # Detect "City, ST" prefix.
            loc_m = re.match(r"^([A-Z][A-Za-z .'-]+,\s*[A-Z]{2})\s*\|\s*(.*)$", first_rest)
            if loc_m:
                location = loc_m.group(1)
                rest[0] = loc_m.group(2)
            else:
                location = first_rest
                rest = rest[1:]
        contact = " | ".join(p for p in rest if p)
    else:
        # Compact format: header_lines after name = [title, location, contact] in that order.
        if len(contact_block) > 3:
            location = contact_block[2]
            contact = contact_block[3]
        elif len(contact_block) > 2:
            third = contact_block[2]
            if "@" in third or "|" in third:
                contact = third
            else:
                location = third

    # Split body into sections by ALL-CAPS headers
    sections: dict[str, str] = {}
    current_section: str | None = None
    current_lines: list[str] = []

    for line in lines[body_start:]:
        stripped = line.strip()
        # Detect section headers: all-caps line, \u22653 chars, has at least one
        # letter (otherwise "2010 - 2011" passes the upper() check), no
        # leading bullet markers.
        has_letter = any(c.isalpha() for c in stripped)
        if (
            stripped
            and has_letter
            and stripped == stripped.upper()
            and not stripped.startswith(("-", "*", "\u2022"))
            and len(stripped) > 3
        ):
            if current_section:
                sections[current_section] = "\n".join(current_lines).strip()
            current_section = stripped
            current_lines = []
        else:
            current_lines.append(line)

    if current_section:
        sections[current_section] = "\n".join(current_lines).strip()

    return {
        "name": name,
        "title": title,
        "location": location,
        "contact": contact,
        "sections": sections,
    }


def parse_skills(text: str) -> list[tuple[str, str]]:
    """Parse skills section into (category, value) pairs.

    Strips any leading bullet marker (``- ``, ``• ``, ``* ``) so the
    category name comes out clean regardless of whether the source
    used Markdown-style bullets.

    Args:
        text: The TECHNICAL SKILLS section text.

    Returns:
        List of (category_name, skills_string) tuples.
    """
    skills: list[tuple[str, str]] = []
    for line in text.strip().split("\n"):
        line = line.strip()
        if ":" in line:
            cat, val = line.split(":", 1)
            cat = re.sub(r"^[\*\-•]\s*", "", cat).strip()
            skills.append((cat, val.strip()))
    return skills


_BULLET_PREFIXES = ("- ", "\u2022 ", "* ")


def _is_bullet(stripped: str) -> bool:
    """True if the line begins with a recognized bullet marker."""
    return any(stripped.startswith(p) for p in _BULLET_PREFIXES)


def _strip_bullet(stripped: str) -> str:
    """Drop the leading bullet marker from a known-bullet line."""
    for prefix in _BULLET_PREFIXES:
        if stripped.startswith(prefix):
            return stripped[len(prefix):].strip()
    return stripped


def parse_entries(text: str) -> list[dict]:
    """Parse experience/project entries from section text.

    Recognizes any of ``- ``, ``\u2022 ``, or ``* `` as a bullet prefix, so the
    same parser handles both pipeline-tailored resumes (which use ``-``)
    and the hand-authored master (which uses ``*``).

    Captures up to TWO non-bullet lines after the title:
      - ``subtitle`` \u2014 the line right under the title (team / department,
        or for legacy entries the date range itself).
      - ``meta``    \u2014 an optional third line for ``date | location`` data
        when the entry uses the ATS-friendly 3-line header format.

    Args:
        text: The EXPERIENCE or PROJECTS section text.

    Returns:
        List of ``{"title": str, "subtitle": str, "meta": str,
                   "bullets": list[str]}`` dicts.
    """
    entries: list[dict] = []
    lines = text.strip().split("\n")
    current: dict | None = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if _is_bullet(stripped):
            if current:
                current["bullets"].append(_strip_bullet(stripped))
        elif current is None or len(current.get("bullets", [])) > 0:
            # New entry: either this is the first row, or we've already seen
            # at least one bullet for the previous entry (so the next bare
            # line is a new job/project header).
            if current:
                entries.append(current)
            current = {"title": stripped, "subtitle": "", "meta": "", "bullets": []}
        elif current and not current["subtitle"]:
            current["subtitle"] = stripped
        elif current and not current.get("meta"):
            current["meta"] = stripped
        else:
            if current:
                current["bullets"].append(stripped)

    if current:
        entries.append(current)

    return entries


# \u2500\u2500 Date / metadata-line detection \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500

# Catches "February 2024 - May 2025", "2024-2025", "Jan 2024 \u2013 Present",
# optionally followed by "  |  Location, ST". Used to separate a team
# subtitle from a date+location meta line.
_DATE_HEAD_RE = re.compile(
    r"""
    ^\s*
    (?:[A-Za-z]+\s+)?\d{4}             # "February 2024" or just "2024"
    \s*[-\u2013\u2014]\s*
    (?:Present|Current|(?:[A-Za-z]+\s+)?\d{4})
    """,
    re.VERBOSE | re.IGNORECASE,
)


def _looks_like_date(s: str) -> bool:
    return bool(s and _DATE_HEAD_RE.match(s))


def _resolve_entry_lines(entry: dict) -> tuple[str, str, str]:
    """Return (title, team_subtitle, dates_and_location).

    Handles both formats:
      * Legacy 2-line: title + (date-only subtitle).
      * New 3-line:    title + team subtitle + date+location meta line.
    """
    title = entry.get("title", "").strip()
    sub = entry.get("subtitle", "").strip()
    meta = entry.get("meta", "").strip()

    if meta and _looks_like_date(meta):
        # Three-line format: subtitle is the team, meta is date+location.
        return title, sub, meta
    if _looks_like_date(sub):
        # Legacy two-line format: subtitle is the date itself.
        return title, "", sub
    return title, sub, meta


# ── Education parsing ────────────────────────────────────────────────────

# Just a year range, optionally wrapped in parens — used to recognize the
# date line in 3-line education entries.
_EDU_DATE_RE = re.compile(
    r"""
    ^\s*\(?\s*
    (?P<start>\d{4})
    \s*[-–—]\s*
    (?P<end>\d{4}|Present|Current)
    \s*\)?\s*$
    """,
    re.VERBOSE | re.IGNORECASE,
)


def parse_education_entries(text: str) -> list[dict]:
    """Parse the EDUCATION section into structured entries.

    Each entry is a blank-line-separated block of 1–3 lines::

        Institution
        [studyType, area  |  area-only]
        [YYYY - YYYY  |  YYYY - Present]

    Legacy single-line ``Field | Institution | Certifications: ...`` rows
    parse as a fall-back (used by older tailored resumes).
    """
    entries: list[dict] = []
    current: dict | None = None

    for raw in text.strip().split("\n"):
        line = raw.strip()
        if not line:
            if current:
                entries.append(current)
                current = None
            continue

        # Legacy "Field | Institution | Certifications: ..." line — keep
        # working for the older pipeline-tailored format.
        if "|" in line:
            if current:
                entries.append(current)
            entries.append({"_legacy_line": line})
            current = None
            continue

        # Year-range line attaches to the current entry.
        if current is not None and (m := _EDU_DATE_RE.match(line)):
            current["startDate"] = m.group("start")
            end = m.group("end")
            if end.lower() not in ("present", "current"):
                current["endDate"] = end
            continue

        if current is None:
            current = {"institution": line}
            continue

        # Second non-date line: either "Degree, Field" or just "Field".
        if "studyType" not in current and "area" not in current:
            if "," in line:
                study, _, area = line.partition(",")
                current["studyType"] = study.strip()
                current["area"] = area.strip()
            else:
                current["area"] = line
            continue

        # Anything beyond that — start a fresh entry (no blank-line break).
        entries.append(current)
        current = {"institution": line}

    if current:
        entries.append(current)

    return entries


# ── HTML Template ────────────────────────────────────────────────────────

def build_html(resume: dict) -> str:
    """Build professional resume HTML from parsed data.

    Args:
        resume: Parsed resume dict from parse_resume().

    Returns:
        Complete HTML string ready for PDF rendering.
    """
    from html import escape as _esc

    def _linkify(text: str) -> str:
        """HTML-escape ``text`` and convert URLs/emails to <a> tags."""
        out = []
        for segment, href in _split_text_with_links(text):
            if href:
                out.append(f'<a href="{_esc(href, quote=True)}">{_esc(segment)}</a>')
            else:
                out.append(_esc(segment))
        return "".join(out)

    sections = resume["sections"]

    # Skills
    skills_html = ""
    if "TECHNICAL SKILLS" in sections:
        skills = parse_skills(sections["TECHNICAL SKILLS"])
        rows = ""
        for cat, val in skills:
            rows += f'<div class="skill-row"><span class="skill-cat">{_esc(cat)}:</span> {_linkify(val)}</div>\n'
        skills_html = f'<div class="section"><div class="section-title">Technical Skills</div>{rows}</div>'

    def _render_entries_html(section_text: str) -> str:
        """Render entries with the 3-line ATS header (title / team / date)."""
        items = ""
        for e in parse_entries(section_text):
            title, team_subtitle, date_meta = _resolve_entry_lines(e)
            bullets = "".join(f"<li>{_linkify(b)}</li>" for b in e["bullets"])
            subtitle_html = (
                f'<div class="entry-subtitle">{_linkify(team_subtitle)}</div>'
                if team_subtitle else ""
            )
            meta_html = (
                f'<div class="entry-meta">{_linkify(date_meta)}</div>'
                if date_meta else ""
            )
            items += (
                f'<div class="entry">'
                f'<div class="entry-title">{_linkify(title)}</div>'
                f'{subtitle_html}{meta_html}'
                f'<ul>{bullets}</ul>'
                f'</div>'
            )
        return items

    # Professional Experience — accept either "PROFESSIONAL EXPERIENCE"
    # (master) or "EXPERIENCE" (pipeline-tailored).
    exp_html = ""
    exp_section = sections.get("PROFESSIONAL EXPERIENCE") or sections.get("EXPERIENCE")
    if exp_section:
        items = _render_entries_html(exp_section)
        exp_html = (
            f'<div class="section">'
            f'<div class="section-title">Professional Experience</div>{items}</div>'
        )

    # Earlier Experience
    earlier_html = ""
    if "EARLIER EXPERIENCE" in sections:
        items = _render_entries_html(sections["EARLIER EXPERIENCE"])
        earlier_html = (
            f'<div class="section">'
            f'<div class="section-title">Earlier Experience</div>{items}</div>'
        )

    # Projects
    proj_html = ""
    if "PROJECTS" in sections:
        items = _render_entries_html(sections["PROJECTS"])
        proj_html = f'<div class="section"><div class="section-title">Projects</div>{items}</div>'

    # Education — rich (institution / degree / dates) blocks.
    def _education_html(text: str) -> str:
        rows = ""
        for entry in parse_education_entries(text):
            if entry.get("_legacy_line"):
                rows += f'<div class="list-row">{_linkify(entry["_legacy_line"])}</div>'
                continue
            inst = entry.get("institution", "")
            degree_bits = [b for b in (entry.get("studyType"), entry.get("area")) if b]
            degree_line = ", ".join(degree_bits)
            dates = ""
            if entry.get("startDate"):
                end = entry.get("endDate") or "Present"
                dates = f"{entry['startDate']} - {end}"
            rows += '<div class="edu-entry">'
            if inst:
                rows += f'<div class="edu-institution">{_linkify(inst)}</div>'
            if degree_line:
                rows += f'<div class="edu-degree">{_linkify(degree_line)}</div>'
            if dates:
                rows += f'<div class="edu-dates">{dates}</div>'
            rows += '</div>'
        return rows

    # Certifications / Languages — flat per-line sections.
    def _list_section_html(label: str, text: str) -> str:
        rows = ""
        for line in text.strip().split("\n"):
            line = line.strip()
            if line:
                rows += f'<div class="list-row">{_linkify(line)}</div>'
        return f'<div class="section"><div class="section-title">{label}</div>{rows}</div>' if rows else ""

    edu_html = (
        f'<div class="section"><div class="section-title">Education</div>{_education_html(sections["EDUCATION"])}</div>'
        if "EDUCATION" in sections else ""
    )
    cert_html = _list_section_html("Certifications", sections["CERTIFICATIONS"]) if "CERTIFICATIONS" in sections else ""
    lang_html = _list_section_html("Languages", sections["LANGUAGES"]) if "LANGUAGES" in sections else ""

    # Summary
    summary_html = ""
    if "SUMMARY" in sections:
        summary_html = (
            f'<div class="section"><div class="section-title">Summary</div>'
            f'<div class="summary">{_linkify(sections["SUMMARY"].strip())}</div></div>'
        )

    # Contact line parsing — keep the visual " | " separators but linkify each part.
    contact = resume["contact"]
    contact_parts = [_linkify(p.strip()) for p in contact.split("|")] if contact else []
    contact_html = " &nbsp;|&nbsp; ".join(contact_parts)

    # Location line (may be empty)
    location_html = f'<div class="location">{_esc(resume["location"])}</div>' if resume["location"] else ""

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
@page {{
    size: letter;
    margin: 0.35in 0.5in;
}}
* {{
    margin: 0;
    padding: 0;
    box-sizing: border-box;
}}
body {{
    font-family: 'Calibri', 'Segoe UI', Arial, sans-serif;
    font-size: 10pt;
    line-height: 1.35;
    color: #1a1a1a;
}}
.header {{
    text-align: center;
    margin-bottom: 4px;
    padding-bottom: 4px;
    border-bottom: 1.5px solid #2a7ab5;
}}
.name {{
    font-size: 18pt;
    font-weight: 700;
    color: #1a3a5c;
    letter-spacing: 0.5px;
}}
.title {{
    font-size: 10.5pt;
    color: #3a6b8c;
    margin: 1px 0;
}}
.location {{
    font-size: 9pt;
    color: #555;
}}
.contact {{
    font-size: 9pt;
    color: #444;
    margin-top: 1px;
}}
/* All hyperlinks: brand color, no underline, never visited-styled. */
a, a:link, a:visited, a:hover, a:active {{
    color: #2a7ab5;
    text-decoration: none;
}}
.contact a {{
    color: #2c3e50;
}}
.section {{
    margin-top: 5px;
}}
.section-title {{
    font-size: 10pt;
    font-weight: 700;
    color: #1a3a5c;
    text-transform: uppercase;
    letter-spacing: 0.8px;
    border-bottom: 1.5px solid #2a7ab5;
    padding-bottom: 1px;
    margin-bottom: 3px;
}}
.summary {{
    font-size: 9.5pt;
    color: #333;
    line-height: 1.4;
}}
.skill-row {{
    font-size: 9.5pt;
    margin: 0;
    line-height: 1.35;
}}
.skill-cat {{
    font-weight: 600;
    color: #1a3a5c;
}}
.entry {{
    margin-bottom: 4px;
    break-inside: avoid;
}}
.entry-title {{
    font-weight: 600;
    font-size: 10pt;
    color: #1a3a5c;
}}
.entry-subtitle {{
    font-size: 9pt;
    color: #4a7a9b;
    font-style: italic;
    margin-bottom: 0;
}}
.entry-meta {{
    font-size: 9pt;
    color: #555;
    margin-bottom: 1px;
}}
ul {{
    margin-left: 14px;
    padding: 0;
}}
li {{
    font-size: 9.5pt;
    margin-bottom: 1px;
    line-height: 1.35;
}}
.list-row {{
    font-size: 9.5pt;
    margin: 0;
    line-height: 1.35;
}}
.edu {{
    font-size: 10pt;
}}
.edu-entry {{
    margin-bottom: 3px;
}}
.edu-institution {{
    font-weight: 600;
    font-size: 10pt;
    color: #1a3a5c;
}}
.edu-degree {{
    font-size: 9.5pt;
    color: #1a1a1a;
}}
.edu-dates {{
    font-size: 9pt;
    color: #555;
}}
</style>
</head>
<body>
<div class="header">
    <div class="name">{_esc(resume['name'])}</div>
    <div class="title">{_esc(resume['title'])}</div>
    {location_html}
    <div class="contact">{contact_html}</div>
</div>
{summary_html}
{skills_html}
{exp_html}
{earlier_html}
{proj_html}
{edu_html}
{cert_html}
{lang_html}
</body>
</html>"""


# ── PDF Renderer ─────────────────────────────────────────────────────────

def render_pdf(html: str, output_path: str, metadata: dict | None = None) -> None:
    """Render HTML to PDF using Playwright's headless Chromium.

    Chromium leaves the PDF's Info dict mostly empty (no Title/Author/etc.)
    and stamps Creator="Chromium". After rendering we post-process the file
    with pypdf to populate the Info dict from ``metadata`` so reviewers,
    ATS parsers, and search indexers see the same fields the DOCX advertises.

    Args:
        html: Complete HTML string.
        output_path: Path to write the PDF file.
        metadata: Optional dict matching the DOCX metadata schema —
            keys: 'title', 'subject', 'author', 'keywords' (str or list),
            'description'. None leaves the Info dict at Chromium's default.
    """
    from patchright.sync_api import sync_playwright

    with sync_playwright() as p:
        # chromium_sandbox=True drops patchright's default --no-sandbox flag.
        browser = p.chromium.launch(chromium_sandbox=True)
        page = browser.new_page()
        page.set_content(html, wait_until="networkidle")
        page.pdf(
            path=output_path,
            format="Letter",
            margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
            print_background=True,
        )
        browser.close()

    if metadata:
        _set_pdf_metadata(output_path, metadata)


def _set_pdf_metadata(path: str, metadata: dict) -> None:
    """Write ``metadata`` into the PDF's Info dictionary (in place).

    Chromium's headless PDF lacks Title/Author/Subject/Keywords. Stamp them
    here so the file looks like a Word-exported PDF instead of a printed
    web page. Also overrides Creator from "Chromium" to a neutral value.
    """
    from pypdf import PdfReader, PdfWriter
    from datetime import datetime, timezone as _tz

    reader = PdfReader(path)
    writer = PdfWriter(clone_from=reader)

    kw = metadata.get("keywords")
    if isinstance(kw, list):
        kw = ", ".join(str(k).strip() for k in kw if k)

    now = datetime.now(_tz.utc)
    pdf_date = now.strftime("D:%Y%m%d%H%M%S+00'00'")

    info = {
        "/Title":        str(metadata.get("title", ""))[:512],
        "/Author":       str(metadata.get("author", "")),
        "/Subject":      str(metadata.get("subject", ""))[:512],
        "/Keywords":     str(kw or ""),
        "/Creator":      "Microsoft Word",
        "/Producer":     "Microsoft Word",
        "/CreationDate": pdf_date,
        "/ModDate":      pdf_date,
    }
    # Drop empties so we don't write zero-length fields.
    info = {k: v for k, v in info.items() if v}
    writer.add_metadata(info)

    with open(path, "wb") as f:
        writer.write(f)


# ── DOCX Renderer ────────────────────────────────────────────────────

def render_docx(resume: dict, output_path: str, metadata: dict | None = None) -> None:
    """Render parsed resume data to a DOCX file using python-docx.

    Args:
        resume: Parsed resume dict from parse_resume().
        output_path: Path to write the DOCX file.
        metadata: Optional dict to populate the DOCX core_properties. Supported
            keys: 'title', 'subject', 'keywords' (str or list), 'description',
            'author', 'company', 'category', 'comments'.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Brand link color (matches the section-heading bottom-border).
    LINK_COLOR_HEX = "2A7AB5"

    def _add_hyperlink_run(paragraph, display: str, href: str,
                           font_size: Pt | None = None,
                           bold: bool = False, italic: bool = False) -> None:
        """Append a clickable hyperlink run to ``paragraph``.

        Renders without an underline (Jobscan-friendly + matches the user's
        styling preference). Falls back to a plain styled run on failure.
        """
        try:
            r_id = paragraph.part.relate_to(
                href,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )
        except Exception:
            r = paragraph.add_run(display)
            r.font.color.rgb = RGBColor(0x2A, 0x7A, 0xB5)
            if font_size:
                r.font.size = font_size
            if bold:
                r.bold = True
            if italic:
                r.italic = True
            return

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        color = OxmlElement("w:color")
        color.set(qn("w:val"), LINK_COLOR_HEX)
        rPr.append(color)

        # Explicitly disable underline (LO/Word default for hyperlink style).
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "none")
        rPr.append(u)

        if font_size is not None:
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(int(font_size.pt * 2)))  # half-points
            rPr.append(sz)
        if bold:
            rPr.append(OxmlElement("w:b"))
        if italic:
            rPr.append(OxmlElement("w:i"))

        # Intentionally NOT setting w:rStyle="Hyperlink" — that style carries
        # a forced underline. We let the explicit w:color + w:u="none" win.
        new_run.append(rPr)
        text_el = OxmlElement("w:t")
        text_el.text = display
        text_el.set(qn("xml:space"), "preserve")
        new_run.append(text_el)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _add_runs_with_links(paragraph, text: str, *,
                             font_size: Pt | None = None,
                             color: RGBColor | None = None,
                             bold: bool = False, italic: bool = False) -> None:
        """Add ``text`` to ``paragraph`` with auto-detected URLs as hyperlinks.

        Plain segments inherit the supplied font_size/color/bold/italic.
        Hyperlinks override color to the brand link color and never underline.
        """
        for segment, href in _split_text_with_links(text):
            if href:
                _add_hyperlink_run(paragraph, segment, href,
                                   font_size=font_size, bold=bold, italic=italic)
            else:
                r = paragraph.add_run(segment)
                if font_size is not None:
                    r.font.size = font_size
                if color is not None:
                    r.font.color.rgb = color
                if bold:
                    r.bold = True
                if italic:
                    r.italic = True

    doc = Document()

    # Page margins (match PDF: 0.35in top/bottom, 0.5in left/right)
    for section in doc.sections:
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # ── Word built-in style customization ──────────────────────────────
    # ATS parsers (Workday/Greenhouse/etc.) walk the OOXML style outline
    # to identify section structure. We use the canonical Title /
    # Subtitle / Heading 1 / Heading 2 styles and re-skin them to match
    # the visual design — same look, but now there's a structural map.

    def _restyle(name: str, *, size: float, bold: bool = False,
                 italic: bool = False, color: tuple[int, int, int] = (0x1A, 0x1A, 0x1A),
                 caps: bool = False,
                 space_before: float = 0, space_after: float = 0,
                 line_spacing: float = 1.35) -> Any:
        s = doc.styles[name]
        f = s.font
        f.name = "Calibri"
        f.size = Pt(size)
        f.bold = bold
        f.italic = italic
        f.color.rgb = RGBColor(*color)
        f.all_caps = caps
        pf = s.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = line_spacing
        # Strip the standard "based on Heading X" parent-style coupling so
        # our overrides aren't trampled by Word's defaults.
        return s

    _restyle("Normal", size=10, color=(0x1A, 0x1A, 0x1A))

    _restyle(
        "Title", size=18, bold=True, color=(0x1A, 0x3A, 0x5C),
        space_after=1, line_spacing=1.1,
    )
    _restyle(
        "Subtitle", size=10.5, color=(0x3A, 0x6B, 0x8C),
        space_after=0, line_spacing=1.2,
    )
    _restyle(
        "Heading 1", size=10, bold=True, color=(0x1A, 0x3A, 0x5C), caps=True,
        space_before=6, space_after=2,
    )
    _restyle(
        "Heading 2", size=10, bold=True, color=(0x1A, 0x3A, 0x5C),
        space_before=4, space_after=0,
    )
    _restyle(
        "Heading 3", size=9, italic=True, color=(0x4A, 0x7A, 0x9B),
        space_after=0,
    )

    # ── Header (Title + Subtitle + Normal) ─────────────────────────────
    name_para = doc.add_paragraph(style="Title")
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.add_run(resume["name"])

    if resume["title"]:
        sub_para = doc.add_paragraph(style="Subtitle")
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_para.add_run(resume["title"])

    if resume["location"]:
        loc_para = doc.add_paragraph()
        loc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        loc_run = loc_para.add_run(resume["location"])
        loc_run.font.size = Pt(9)
        loc_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if resume["contact"]:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_runs_with_links(
            contact_para, resume["contact"],
            font_size=Pt(9), color=RGBColor(0x44, 0x44, 0x44),
        )

    sections = resume["sections"]

    def _add_section_heading(title: str) -> None:
        """Section heading — Heading 1 style, plus a bottom border."""
        para = doc.add_paragraph(style="Heading 1")
        para.add_run(title.upper())
        # Bottom border via XML (matches the PDF blue underline).
        pPr = para._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "2A7AB5",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _render_entries(section_text: str) -> None:
        """Render PROFESSIONAL EXPERIENCE / PROJECTS / EARLIER EXPERIENCE
        entries with the 3-line ATS-friendly header (title / team / dates).
        """
        for entry in parse_entries(section_text):
            title, team_subtitle, date_meta = _resolve_entry_lines(entry)

            # Entry title — Heading 2 (gives ATS a clear company/role marker).
            tp = doc.add_paragraph(style="Heading 2")
            _add_runs_with_links(tp, title, font_size=Pt(10), bold=True)

            if team_subtitle:
                sp = doc.add_paragraph(style="Heading 3")
                _add_runs_with_links(
                    sp, team_subtitle,
                    font_size=Pt(9), color=RGBColor(0x4A, 0x7A, 0x9B), italic=True,
                )

            if date_meta:
                dp = doc.add_paragraph()
                _add_runs_with_links(
                    dp, date_meta,
                    font_size=Pt(9), color=RGBColor(0x55, 0x55, 0x55),
                )

            for bullet in entry["bullets"]:
                bp = doc.add_paragraph(style="List Bullet")
                _add_runs_with_links(bp, bullet, font_size=Pt(9.5))

    # ── Summary ────────────────────────────────────────────────────────
    if "SUMMARY" in sections:
        _add_section_heading("Summary")
        p = doc.add_paragraph()
        _add_runs_with_links(
            p, sections["SUMMARY"].strip(),
            font_size=Pt(9.5), color=RGBColor(0x33, 0x33, 0x33),
        )

    # ── Technical Skills ───────────────────────────────────────────────
    if "TECHNICAL SKILLS" in sections:
        _add_section_heading("Technical Skills")
        for cat, val in parse_skills(sections["TECHNICAL SKILLS"]):
            p = doc.add_paragraph()
            cat_run = p.add_run(f"{cat}: ")
            cat_run.bold = True
            cat_run.font.size = Pt(9.5)
            cat_run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
            _add_runs_with_links(p, val, font_size=Pt(9.5))

    # ── Professional Experience (Jobscan §4: prefer the full label) ──
    # Pipeline-tailored resumes use "EXPERIENCE"; the master uses
    # "PROFESSIONAL EXPERIENCE". Either lands here.
    exp_section = sections.get("PROFESSIONAL EXPERIENCE") or sections.get("EXPERIENCE")
    if exp_section:
        _add_section_heading("Professional Experience")
        _render_entries(exp_section)

    # ── Earlier Experience ─────────────────────────────────────────────
    if "EARLIER EXPERIENCE" in sections:
        _add_section_heading("Earlier Experience")
        _render_entries(sections["EARLIER EXPERIENCE"])

    # ── Projects ───────────────────────────────────────────────────────
    if "PROJECTS" in sections:
        _add_section_heading("Projects")
        _render_entries(sections["PROJECTS"])

    # ── Education ──────────────────────────────────────────────────────
    if "EDUCATION" in sections:
        _add_section_heading("Education")
        for entry in parse_education_entries(sections["EDUCATION"]):
            if entry.get("_legacy_line"):
                p = doc.add_paragraph()
                _add_runs_with_links(p, entry["_legacy_line"], font_size=Pt(10))
                continue

            # Institution: bold (Heading 2-ish weight without the spacing).
            if entry.get("institution"):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                _add_runs_with_links(
                    p, entry["institution"],
                    font_size=Pt(10), color=RGBColor(0x1A, 0x3A, 0x5C), bold=True,
                )

            # Degree + field on one line.
            degree_bits = [b for b in (entry.get("studyType"), entry.get("area")) if b]
            if degree_bits:
                p = doc.add_paragraph()
                _add_runs_with_links(p, ", ".join(degree_bits), font_size=Pt(9.5))

            # Date range — small gray under the entry.
            if entry.get("startDate"):
                end = entry.get("endDate") or "Present"
                p = doc.add_paragraph()
                _add_runs_with_links(
                    p, f"{entry['startDate']} - {end}",
                    font_size=Pt(9), color=RGBColor(0x55, 0x55, 0x55),
                )

    # ── Certifications ─────────────────────────────────────────────────
    if "CERTIFICATIONS" in sections:
        _add_section_heading("Certifications")
        for line in sections["CERTIFICATIONS"].strip().split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            _add_runs_with_links(p, line, font_size=Pt(9.5))

    # ── Languages ──────────────────────────────────────────────────────
    if "LANGUAGES" in sections:
        _add_section_heading("Languages")
        for line in sections["LANGUAGES"].strip().split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            _add_runs_with_links(p, line, font_size=Pt(9.5))

    # Populate core_properties. Always overwrite python-docx's template
    # defaults (author='python-docx', comments='generated by python-docx',
    # created/modified=2013-12-23) so reviewers + ATS systems don't see
    # tooling fingerprints.
    cp = doc.core_properties
    from datetime import datetime, timezone as _tz
    now = datetime.now(_tz.utc).replace(microsecond=0)

    metadata = metadata or {}
    cp.title    = str(metadata.get("title", ""))[:256]
    cp.subject  = str(metadata.get("subject", ""))[:256]
    cp.author   = str(metadata.get("author", "")) or ""
    cp.company  = str(metadata.get("company", "")) or ""
    cp.category = str(metadata.get("category", "")) or ""
    cp.last_modified_by = str(metadata.get("last_modified_by", metadata.get("author", "")))[:256]
    # `comments` literally inherits "generated by python-docx" from the
    # template — overwrite (with description if provided, else blank).
    cp.comments = str(metadata.get("comments", metadata.get("description", "")))[:2000]
    # Refresh timestamps. created keeps any caller-supplied value (so the
    # master can preserve its real authored date) but must not stay at the
    # 2013 template default.
    created = metadata.get("created")
    cp.created = created if isinstance(created, datetime) else now
    cp.modified = now

    kw = metadata.get("keywords")
    if kw:
        if isinstance(kw, list):
            kw = ", ".join(str(k).strip() for k in kw if k)
        # OOXML core-properties hard-caps `keywords` at 255 chars; the
        # python-docx setter raises ValueError beyond that.
        cp.keywords = str(kw)[:255]
    else:
        cp.keywords = ""

    # Save once so app.xml + the package are written, then patch the
    # extended-properties to clobber python-docx's "Microsoft Macintosh
    # Word" default Application string.
    doc.save(output_path)
    _scrub_docx_app_xml(output_path)


def _scrub_docx_app_xml(path: str) -> None:
    """Replace python-docx template's app.xml fingerprints in place.

    python-docx ships an app.xml that sets ``<Application>Microsoft
    Macintosh Word</Application>`` regardless of the host system. That
    string is a known python-docx tell. Rewrite the file's app.xml so
    the Application string reads ``Microsoft Office Word``, which is
    what an actual Word save produces.
    """
    import zipfile
    import shutil
    import tempfile
    import os
    src = str(path)
    fd, tmp = tempfile.mkstemp(suffix=".docx", dir=os.path.dirname(src) or None)
    os.close(fd)
    try:
        with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "docProps/app.xml":
                    data = data.replace(
                        b"<Application>Microsoft Macintosh Word</Application>",
                        b"<Application>Microsoft Office Word</Application>",
                    )
                zout.writestr(item, data)
        shutil.move(tmp, src)
    except Exception:
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise


# ── Public API ───────────────────────────────────────────────────────────

def convert_to_pdf(
    text_path: Path,
    output_path: Path | None = None,
    html_only: bool = False,
    doc_format: str = "docx",
    metadata: dict | None = None,
) -> Path:
    """Convert a text resume/cover letter to PDF or DOCX.

    Args:
        text_path: Path to the .txt file to convert.
        output_path: Optional override for the output path. Defaults to same
            name with the appropriate extension.
        html_only: If True, output HTML instead of PDF/DOCX.
        doc_format: Output format — "docx" (default) or "pdf".
        metadata: Optional dict to populate DOCX core_properties (ignored for PDF).
            Supported keys: 'title', 'subject', 'keywords', 'author', 'company',
            'category', 'comments', 'description'.

    Returns:
        Path to the generated file.
    """
    if doc_format not in VALID_DOC_FORMATS:
        raise ValueError(f"Invalid doc_format '{doc_format}'. Must be one of: {VALID_DOC_FORMATS}")

    text_path = Path(text_path)
    text = text_path.read_text(encoding="utf-8")
    resume = parse_resume(text)

    if html_only:
        html = build_html(resume)
        out = output_path or text_path.with_suffix(".html")
        out = Path(out)
        out.write_text(html, encoding="utf-8")
        log.info("HTML generated: %s", out)
        return out

    if doc_format == "docx":
        out = output_path or text_path.with_suffix(".docx")
        out = Path(out)
        render_docx(resume, str(out), metadata=metadata)
        log.info("DOCX generated: %s", out)
        return out

    # Default: PDF
    html = build_html(resume)
    out = output_path or text_path.with_suffix(".pdf")
    out = Path(out)
    render_pdf(html, str(out), metadata=metadata)
    log.info("PDF generated: %s", out)
    return out


def batch_convert(limit: int = 50, doc_format: str = "docx") -> int:
    """Convert .txt files in TAILORED_DIR that don't have corresponding output files.

    Scans for .txt files (excluding _JOB.txt and _REPORT.json), checks if a
    file with the target extension already exists, and converts any that are missing.

    Args:
        limit: Maximum number of files to convert.
        doc_format: Output format — "docx" (default) or "pdf".

    Returns:
        Number of files generated.
    """
    if doc_format not in VALID_DOC_FORMATS:
        raise ValueError(f"Invalid doc_format '{doc_format}'. Must be one of: {VALID_DOC_FORMATS}")

    ext = f".{doc_format}"

    if not TAILORED_DIR.exists():
        log.warning("Tailored directory does not exist: %s", TAILORED_DIR)
        return 0

    txt_files = sorted(TAILORED_DIR.glob("*.txt"))
    # Exclude _JOB.txt files from resume conversion
    # (they get their own conversion calls)
    candidates = [
        f for f in txt_files
        if not f.name.endswith("_JOB.txt")
    ]

    # Filter to those without a corresponding output file
    to_convert: list[Path] = []
    for f in candidates:
        out_path = f.with_suffix(ext)
        if not out_path.exists():
            to_convert.append(f)
        if len(to_convert) >= limit:
            break

    if not to_convert:
        log.debug("All text files already have %s files.", doc_format.upper())
        return 0

    log.info("Converting %d files to %s...", len(to_convert), doc_format.upper())
    converted = 0
    for f in to_convert:
        try:
            convert_to_pdf(f, doc_format=doc_format)
            converted += 1
        except Exception as e:
            log.error("Failed to convert %s: %s", f.name, e)

    log.info("Done: %d/%d %s files generated in %s", converted, len(to_convert), doc_format.upper(), TAILORED_DIR)
    return converted
